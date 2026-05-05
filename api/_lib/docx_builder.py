from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

GEORGIA = "Georgia"
COURIER = "Courier New"
RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x0A, 0x0A, 0x0F)
DIAGRAM_GRAY = RGBColor(0x88, 0x88, 0x88)

BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3"]


def _apply_run_formatting(
    run: Any,
    *,
    font_name: str,
    size_pt: float,
    bold: bool,
    italic: bool,
    color: Any = None,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_title(document: Any, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title)
    _apply_run_formatting(run, font_name=GEORGIA, size_pt=24, bold=True, italic=False, color=BLACK)


def _add_heading(document: Any, section: Dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    size = 18 if int(section.get("headingLevel", 1)) <= 1 else 14
    color = RED if section.get("headingColor") == "red" else BLACK
    run = paragraph.add_run(section.get("heading", ""))
    _apply_run_formatting(run, font_name=GEORGIA, size_pt=size, bold=True, italic=False, color=color)


def _paragraph_block(document: Any, block: Dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    indent_level = int(block.get("indentLevel", 0))
    if indent_level > 0:
        paragraph.paragraph_format.left_indent = Inches(indent_level * 0.5)
    run = paragraph.add_run(block.get("text", ""))
    _apply_run_formatting(
        run,
        font_name=GEORGIA,
        size_pt=11,
        bold=bool(block.get("isBold")),
        italic=bool(block.get("isItalic")),
        color=BLACK,
    )


def _bullet_block(document: Any, block: Dict[str, Any]) -> None:
    indent_level = int(block.get("indentLevel", 0))
    style = BULLET_STYLES[min(indent_level, len(BULLET_STYLES) - 1)]
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    if indent_level >= len(BULLET_STYLES):
        extra = (indent_level - (len(BULLET_STYLES) - 1)) * 0.5
        paragraph.paragraph_format.left_indent = Inches((len(BULLET_STYLES) - 1) * 0.5 + extra)
    run = paragraph.add_run(block.get("text", ""))
    _apply_run_formatting(
        run,
        font_name=GEORGIA,
        size_pt=11,
        bold=bool(block.get("isBold")),
        italic=bool(block.get("isItalic")),
        color=BLACK,
    )


def _numbered_block(document: Any, block: Dict[str, Any]) -> None:
    indent_level = int(block.get("indentLevel", 0))
    style = NUMBER_STYLES[min(indent_level, len(NUMBER_STYLES) - 1)]
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    if indent_level >= len(NUMBER_STYLES):
        extra = (indent_level - (len(NUMBER_STYLES) - 1)) * 0.5
        paragraph.paragraph_format.left_indent = Inches((len(NUMBER_STYLES) - 1) * 0.5 + extra)
    run = paragraph.add_run(block.get("text", ""))
    _apply_run_formatting(
        run,
        font_name=GEORGIA,
        size_pt=11,
        bold=bool(block.get("isBold")),
        italic=bool(block.get("isItalic")),
        color=BLACK,
    )


def _equation_block(document: Any, block: Dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(block.get("text", ""))
    _apply_run_formatting(
        run,
        font_name=COURIER,
        size_pt=10,
        bold=bool(block.get("isBold")),
        italic=bool(block.get("isItalic")),
        color=BLACK,
    )


def _diagram_block(document: Any, block: Dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    text = block.get("text", "")
    run = paragraph.add_run(f"[DIAGRAM: {text}]")
    _apply_run_formatting(
        run,
        font_name=GEORGIA,
        size_pt=11,
        bold=False,
        italic=True,
        color=DIAGRAM_GRAY,
    )


_RENDERERS = {
    "paragraph": _paragraph_block,
    "bullet": _bullet_block,
    "numbered": _numbered_block,
    "equation": _equation_block,
    "diagram": _diagram_block,
}


def build_docx_bytes(analysis: Dict[str, Any]) -> bytes:
    document = Document()

    title = analysis.get("title", "")
    if isinstance(title, str) and title.strip():
        _add_title(document, title.strip())

    sections: List[Dict[str, Any]] = analysis.get("sections", []) or []
    for section in sections:
        heading = section.get("heading", "")
        if isinstance(heading, str) and heading.strip():
            _add_heading(document, section)
        content_blocks = section.get("content", []) or []
        for block in content_blocks:
            renderer = _RENDERERS.get(block.get("type", "paragraph"), _paragraph_block)
            renderer(document, block)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
